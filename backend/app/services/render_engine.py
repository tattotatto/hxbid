"""宏曦标书 - Word排版渲染引擎.

Copyright (c) 2026 云南宏曦科技有限公司. All rights reserved.

Converts structured bid chapters into a fully formatted .docx with cover page,
headers, footers, and proper Chinese typography. Also provides PDF export via
LibreOffice headless (best-effort) and a default-template factory.

Handles markdown artifacts that may appear in AI-generated content, converting
them to proper Word formatting (headings, bold, italic, bullets, tables).
"""

import os
import re
import time
import uuid
import io as std_io
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

from app.config import settings

# ── Default style constants (Chinese bid-document spec) ─────────────────

DEFAULT_STYLE = {
    "body_font_name": "宋体",          # 宋体
    "body_font_size": Pt(12),                   # 小四
    "body_line_spacing": 1.5,
    "heading1_font_name": "黑体",       # 黑体 三号
    "heading1_font_size": Pt(16),
    "heading2_font_name": "黑体",       # 黑体 四号
    "heading2_font_size": Pt(14),
    "heading3_font_name": "楷体",       # 楷体 小四号加粗
    "heading3_font_size": Pt(12),
    "margin_top": Cm(2.54),
    "margin_bottom": Cm(2.54),
    "margin_left": Cm(3.17),
    "margin_right": Cm(3.17),
    "header_text": "",
}

# Black colour constant — all text MUST be black in Chinese bid documents
BLACK = RGBColor(0, 0, 0)


# ── Internal helpers ──────────────────────────────────────────────────

def _set_run_font(run, font_name, font_size, bold=False, color=BLACK):
    """Set Western + East-Asian font, size, bold, and colour on a ``Run``."""
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    run.font.color.rgb = color

    # Ensure the w:eastAsia attribute is set for CJK rendering
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _set_paragraph_spacing(paragraph, line_spacing=1.5):
    """Set line spacing on a paragraph (1.5x is the standard for Chinese bids)."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing


def _add_body_paragraph(doc, text, style):
    """Add a standard body paragraph with first-line indent."""
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, style["body_line_spacing"])
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Pt(24)
    run = p.add_run(text)
    _set_run_font(run, style["body_font_name"], style["body_font_size"])
    return p


def _add_body_paragraph_with_bold(doc, segments, style):
    """Add a body paragraph where *segments* is a list of (text, is_bold) tuples.

    Each segment is rendered as a separate run so inline bold works correctly.
    """
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, style["body_line_spacing"])
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Pt(24)
    for text, is_bold in segments:
        if not text:
            continue
        run = p.add_run(text)
        _set_run_font(run, style["body_font_name"], style["body_font_size"], bold=is_bold)
    return p


def _add_page_number(paragraph):
    """Insert a bare PAGE field into *paragraph* using low-level OxmlElement calls.

    Produces the equivalent of Word's ``{ PAGE }`` field so that each page
    footer displays the current page number.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- fldChar begin --
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._element.append(fld_begin)

    # -- instrText --
    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_instr._element.append(instr)

    # -- fldChar end --
    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._element.append(fld_end)


def _add_cover_page(doc, project_name, style, company_name=""):
    """Append a regulated Chinese bid cover page to *doc*."""
    # ── Leading vertical space ──
    for _ in range(6):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)

    # ── Project name (22 pt 黑体 bold, centred) ──
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run(project_name)
    _set_run_font(r_name, style["heading1_font_name"], Pt(22), bold=True)

    doc.add_paragraph()  # blank line

    # ── "投标文件" (18 pt 黑体, centred) ──
    p_bid = doc.add_paragraph()
    p_bid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_bid = p_bid.add_run("投标文件")
    _set_run_font(r_bid, style["heading1_font_name"], Pt(18), bold=False)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Company name (16 pt 黑体 bold, centred) ──
    p_company = doc.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_company = p_company.add_run(company_name or "投标人名称")
    _set_run_font(r_company, style["heading1_font_name"], Pt(16), bold=True)

    doc.add_paragraph()

    # ── Date ──
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run(date.today().strftime("%Y年%m月%d日"))
    _set_run_font(r_date, style["heading1_font_name"], Pt(14), bold=False)

    # Page break so TOC starts on a fresh page
    doc.add_page_break()


def _insert_page_break(doc):
    """Insert a reliable page-break as its own paragraph.

    Uses low-level OxmlElement to create ``<w:br w:type="page"/>`` inside
    an otherwise empty paragraph. This is more reliable than
    ``doc.add_page_break()`` because it guarantees the break element is not
    attached to a preceding content run.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._element.append(br)


def _add_toc_page(doc, chapters, style):
    """Insert a Table of Contents page after the cover page.

    Adds a "目录" heading, then inserts a Word TOC field that auto-populates
    when the document is opened in Microsoft Word. Falls back gracefully in
    other editors (LibreOffice, WPS) which also support TOC fields.
    """
    # ── TOC heading ──
    h = doc.add_heading("目录", level=1)
    for run in h.runs:
        _set_run_font(
            run,
            style["heading1_font_name"],
            style["heading1_font_size"],
            bold=True,
        )

    # ── Insert Word TOC field ──
    # This creates the standard { TOC \o "1-3" \h \z } field that
    # populates from Heading 1, 2, and 3 styles.
    p_toc = doc.add_paragraph()
    p_toc.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # fldChar begin
    r_begin = p_toc.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_begin._element.append(fld_begin)

    # instrText: TOC field code — collect Heading 1-3 entries
    r_instr = p_toc.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z '
    r_instr._element.append(instr)

    # fldChar separate
    r_sep = p_toc.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_sep._element.append(fld_sep)

    # Placeholder text — visible until user updates the TOC in Word
    r_placeholder = p_toc.add_run(
        '（请在Word中右键点击此处，选择"更新域"以生成目录）'
    )
    _set_run_font(r_placeholder, style["body_font_name"], Pt(10), color=RGBColor(128, 128, 128))

    # fldChar end
    r_end = p_toc.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_end._element.append(fld_end)


def _safe_filename(name):
    """Turn *name* into a filesystem-safe slug (no path separators, etc.)."""
    safe = re.sub(r'[<>:"/\\|?*]', "_", name)
    safe = re.sub(r"\s+", "_", safe)
    safe = safe.strip("._")
    return safe if safe else "document"


# ── Markdown detection helpers ────────────────────────────────────────

def _is_table_separator(line):
    """Check if line is a markdown table separator like |---|---| or |:---|:---:|"""
    stripped = line.strip()
    return bool(re.match(r'^\|[\s\-:]+\|', stripped)) and '---' in stripped


def _is_table_row(line):
    """Check if line is a markdown table data row like | col1 | col2 |"""
    stripped = line.strip()
    if not stripped.startswith('|') or not stripped.endswith('|'):
        return False
    # Must have at least 2 pipe-separated cells, and not be a separator
    parts = [c.strip() for c in stripped.split('|')]
    # Remove empty strings from leading/trailing pipes
    parts = [p for p in parts if p]
    return len(parts) >= 2 and not _is_table_separator(line)


def _parse_table_cells(line):
    """Parse | col1 | col2 | col3 | into a list of cell strings."""
    stripped = line.strip()
    # Split by pipe, skip first and last empty elements
    cells = stripped.split('|')
    # Drop leading empty (before first |) and trailing empty (after last |)
    if cells and cells[0] == '':
        cells = cells[1:]
    if cells and cells[-1] == '':
        cells = cells[:-1]
    return [c.strip() for c in cells]


def _is_markdown_heading(line):
    """Check if line is a markdown heading (# / ## / ### )."""
    return bool(re.match(r'^#{1,3}\s+\S', line))


def _get_heading_level_and_text(line):
    """Extract heading level (1-3) and text from a markdown heading line."""
    m = re.match(r'^(#{1,3})\s+(.+)$', line)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        # Remove trailing # marks sometimes left by AI
        text = re.sub(r'\s*#+\s*$', '', text)
        return level, text
    return 0, line


def _is_bullet_item(line):
    """Check if line is a markdown bullet item (- or * at start)."""
    stripped = line.strip()
    return bool(re.match(r'^[\-\*]\s+\S', stripped))


def _get_bullet_text(line):
    """Extract text content from a bullet line."""
    m = re.match(r'^[\-\*]\s+(.+)$', line.strip())
    return m.group(1) if m else line


def _is_numbered_item(line):
    """Check if line is a numbered list item like 1. or 1) or (1) etc."""
    return bool(re.match(r'^\d+[\.\)、]\s*\S', line.strip()))


def _get_numbered_text(line):
    """Extract text from a numbered item, preserving the number."""
    m = re.match(r'^(\d+[\.\)、]\s*)(.+)$', line.strip())
    return m.group(1).strip(), m.group(2).strip() if m else ("", line)


def _is_horizontal_rule(line):
    """Check if line is a markdown horizontal rule (---, ***, ___ )."""
    stripped = line.strip()
    return bool(re.match(r'^[\-\*_]{3,}$', stripped))


# ── Markdown cleanup ──────────────────────────────────────────────────

def _clean_markdown_inline(text):
    """Clean inline markdown from text, returning (segments, had_formatting).

    segments is a list of (text, is_bold) tuples.
    had_formatting is True if any markdown was actually converted.
    """
    segments = []
    had_formatting = False

    # Pattern for **bold** text
    bold_pattern = re.compile(r'\*\*(.+?)\*\*')

    # Split by bold markers, then process italic in each piece
    last_end = 0
    for m in bold_pattern.finditer(text):
        # Add text before this bold segment (process italic within it)
        before = text[last_end:m.start()]
        if before:
            sub_segments, sub_had = _process_italic(before)
            segments.extend(sub_segments)
            if sub_had:
                had_formatting = True

        # Add the bold segment
        bold_text = m.group(1)
        segments.append((bold_text, True))
        had_formatting = True
        last_end = m.end()

    # Remaining text after last bold match
    if last_end < len(text):
        remaining = text[last_end:]
        sub_segments, sub_had = _process_italic(remaining)
        segments.extend(sub_segments)
        if sub_had:
            had_formatting = True

    # If no formatting found, return the original text as-is
    if not had_formatting:
        return [(text, False)], False

    return segments, True


def _process_italic(text):
    """Process *italic* markers within text that has already had **bold** removed."""
    segments = []
    had_formatting = False
    italic_pattern = re.compile(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)')

    last_end = 0
    for m in italic_pattern.finditer(text):
        before = text[last_end:m.start()]
        if before:
            segments.append((before, False))
        italic_text = m.group(1)
        segments.append((italic_text, False))  # italic → normal for bid docs
        had_formatting = True
        last_end = m.end()

    if last_end < len(text):
        segments.append((text[last_end:], False))

    if not had_formatting:
        return [(text, False)], False

    return segments, True


def _clean_lone_symbols(text):
    """Remove orphaned markdown symbols from a line of text.

    Handles: stray * # _ at line boundaries, multiple consecutive symbols.
    """
    # Remove leading/trailing stray asterisks that aren't part of ** pairs
    text = re.sub(r'(?<!\*)\*(?!\*)', '', text)
    # Remove stray # at start of lines that aren't real headings
    text = re.sub(r'^\#{1,2}(?![#\s])', '', text)
    # Remove trailing # marks
    text = re.sub(r'\s*#+\s*$', '', text)
    # Collapse multiple spaces
    text = re.sub(r' {2,}', ' ', text)
    return text.strip()


# ── Semicolon key:value table detection ─────────────────────────────────

def _is_semicolon_table_title(line):
    """Check if line is a table title like 表1：xxx or 表X：xxx"""
    stripped = line.strip()
    return bool(re.match(r'^表[一二三四五六七八九十\d]+[：:]\s*\S', stripped))


def _is_semicolon_table_row(line):
    """Check if line is a semicolon-delimited key:value row.

    Example: 姓名：张三；年龄：36岁；学历：本科；证书：消防员证
    Must have at least 3 semicolons and at least 3 key:value pairs.
    """
    stripped = line.strip()
    if not stripped:
        return False
    # Must contain semicolons with Chinese colon key:value patterns
    # Count "字段：值" pairs separated by semicolons
    parts = stripped.split('；')
    parts = [p.strip() for p in parts if p.strip()]
    if len(parts) < 3:
        return False
    # Each part should contain a Chinese colon separating key:value
    kv_count = sum(1 for p in parts if '：' in p or ':' in p)
    return kv_count >= 3


def _parse_semicolon_table_row(line):
    """Parse a semicolon key:value row into a list of values.

    Example: 姓名：张三；年龄：36岁 → ['张三', '36岁']
    Returns list of cell values (without keys).
    """
    stripped = line.strip()
    parts = stripped.split('；')
    parts = [p.strip() for p in parts if p.strip()]
    values = []
    for part in parts:
        # Split on Chinese colon, take the value part
        if '：' in part:
            _, val = part.split('：', 1)
        elif ':' in part:
            _, val = part.split(':', 1)
        else:
            val = part
        values.append(val.strip())
    return values


def _extract_semicolon_keys(line):
    """Extract column header keys from a semicolon key:value row.

    Example: 姓名：张三；年龄：36岁 → ['姓名', '年龄']
    """
    stripped = line.strip()
    parts = stripped.split('；')
    parts = [p.strip() for p in parts if p.strip()]
    keys = []
    for part in parts:
        if '：' in part:
            key, _ = part.split('：', 1)
        elif ':' in part:
            key, _ = part.split(':', 1)
        else:
            key = ''
        keys.append(key.strip())
    return keys


# ── Table rendering ───────────────────────────────────────────────────

def _render_table(doc, rows, style, is_first_row_header=True):
    """Render a parsed markdown table as a Word table.

    Parameters
    ----------
    doc : Document
    rows : list[list[str]]
        List of rows, each row is a list of cell strings.
    style : dict
        The active document style config.
    is_first_row_header : bool
        Whether the first row should be treated as a header (bold).
    """
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    # Normalise row lengths
    for r in rows:
        while len(r) < num_cols:
            r.append("")

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for ri, row_cells in enumerate(rows):
        for ci, cell_text in enumerate(row_cells):
            cell = table.rows[ri].cells[ci]
            # Clear default empty paragraph
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(cell_text)
            is_header = is_first_row_header and ri == 0
            _set_run_font(
                run,
                style["body_font_name"],
                style["body_font_size"],
                bold=is_header,
            )
            # Cell paragraph formatting
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = cell.paragraphs[0].paragraph_format
            pf.space_before = Pt(2)
            pf.space_after = Pt(2)

    # Add a blank paragraph after the table for spacing
    doc.add_paragraph()


# ── Content preprocessing ─────────────────────────────────────────────

def _preprocess_content(content):
    """Clean up AI-generated content before rendering.

    1. Merge consecutive blank lines
    2. Remove horizontal rules (---, ***)
    3. Remove lines that are purely markdown symbols
    """
    lines = content.split('\n')
    cleaned = []
    blank_count = 0

    for line in lines:
        stripped = line.strip()

        # Skip horizontal rules
        if _is_horizontal_rule(stripped):
            continue

        # Skip lines that are purely markdown symbols
        if re.match(r'^[\#\*\-_]{1,3}$', stripped):
            continue

        # Merge multiple blank lines into one
        if not stripped:
            blank_count += 1
            if blank_count <= 1:
                cleaned.append('')
        else:
            blank_count = 0
            cleaned.append(line)

    return '\n'.join(cleaned)


def _resolve_relative_path(image_path: str) -> Path | None:
    """Resolve a stored image path to an absolute filesystem path.

    Stored paths may or may not include the "uploads/" prefix depending on
    when and how the image was uploaded. Tries multiple resolution strategies.

    Returns the first existing path, or None if the file cannot be found.
    """
    p = Path(image_path)
    if p.is_absolute():
        return p if p.exists() else None

    candidates = [
        Path(settings.UPLOAD_DIR) / image_path,
        Path.cwd() / image_path,
    ]
    # Also try stripping a leading "uploads/" prefix
    s = str(image_path).replace("\\", "/")
    if s.startswith("uploads/"):
        candidates.append(Path(settings.UPLOAD_DIR) / s[len("uploads/"):])

    for c in candidates:
        if c.exists():
            return c
    return None


# ── Image insertion ───────────────────────────────────────────────────

# A4 usable width after 3.17 cm margins on each side ≈ 14.66 cm ≈ 5.77 in
IMAGE_MAX_WIDTH = Cm(14)
IMAGE_MAX_HEIGHT = Cm(20)


def _render_id_card_pair(doc, front_path, front_label, back_path, back_label, style):
    """Render front and back of an ID card side-by-side in a 2-column table."""
    half_width = Cm(6.8)
    fp_front = _resolve_relative_path(front_path)
    fp_back = _resolve_relative_path(back_path)
    if fp_front is None and fp_back is None:
        return
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn("w:val"), "nil")
        tblBorders.append(border_el)
    tblPr.append(tblBorders)
    for col_idx, (fp, label) in enumerate([(fp_front, front_label), (fp_back, back_label)]):
        cell = table.cell(0, col_idx)
        _remove_cell_borders(cell)
        if fp is None:
            continue
        try:
            pil_img = PILImage.open(str(fp))
        except Exception:
            continue
        if pil_img.mode not in ("RGB", "L"):
            pil_img = pil_img.convert("RGB")
        max_px = int(half_width / Cm(1) * 37.795)
        if pil_img.width > max_px:
            ratio = max_px / pil_img.width
            pil_img = pil_img.resize((max_px, int(pil_img.height * ratio)), PILImage.LANCZOS)
        buf = std_io.BytesIO()
        pil_img.save(buf, format="PNG")
        buf.seek(0)
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(buf, width=half_width)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        _set_run_font(r2, style["body_font_name"], Pt(10))
    doc.add_paragraph()


def _insert_image(doc, image_path, label, style, *, no_rotate=False, max_width=None):
    """Insert a single image into the document.

    - Portrait images (height >= width) are inserted as-is.
    - Landscape images (width > height) are rotated 90° clockwise so they
      fit the portrait page orientation, unless *no_rotate* is True.

    The image is scaled to fit within *max_width* (or ``IMAGE_MAX_WIDTH``)
    while preserving aspect ratio.  A centred label is placed below the image.
    """
    if max_width is None:
        max_width = IMAGE_MAX_WIDTH

    full_path = _resolve_relative_path(image_path)
    if full_path is None:
        return  # silently skip — file not found at any candidate path

    # ── Open with PIL to check orientation ──
    try:
        pil_img = PILImage.open(str(full_path))
    except Exception:
        # Silently skip unreadable images
        return
        return

    # Convert RGBA/CMYK/etc. to RGB for JPEG-in-docx compatibility
    if pil_img.mode not in ('RGB', 'L'):
        pil_img = pil_img.convert('RGB')

    is_landscape = pil_img.width > pil_img.height

    if is_landscape and not no_rotate:
        pil_img = pil_img.rotate(270, expand=True)  # 270° CW = portrait

    # ── Scale to fit page width ──
    max_px = int(max_width / Cm(1) * 37.795)  # rough px at ~96 dpi
    if pil_img.width > max_px:
        ratio = max_px / pil_img.width
        new_h = int(pil_img.height * ratio)
        pil_img = pil_img.resize((max_px, new_h), PILImage.LANCZOS)

    # ── Save to in-memory buffer ──
    buf = std_io.BytesIO()
    pil_img.save(buf, format='PNG')
    buf.seek(0)
    width_inches = pil_img.width / 96.0  # approximate

    # ── Insert image into document ──
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(buf, width=Inches(min(width_inches, 5.5)))

    # ── Label below image ──
    p_label = doc.add_paragraph(label)
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = p_label.add_run(label)
    _set_run_font(run_label, style['body_font_name'], Pt(10))
    p_label.paragraph_format.space_after = Pt(12)


def _add_company_footer(doc, style):
    """Append a company information bar at the bottom of the document.

    Layout (3-column borderless table):
      Left:   company logo + company name
      Center: address
      Right:  website
    """
    from app.models.company_profile import CompanyProfile
    from sqlalchemy import create_engine, select
    from sqlalchemy.orm import Session

    # Load company profile from the sync DB
    try:
        sync_url = settings.DATABASE_URL_SYNC
        engine = create_engine(sync_url)
        with Session(engine) as session:
            result = session.execute(select(CompanyProfile).limit(1))
            profile = result.scalar_one_or_none()
    except Exception:
        profile = None

    company_name = profile.company_name if profile else ""
    address = profile.address if profile else ""
    website = profile.website if profile else ""
    logo_path = profile.logo_image if profile else ""

    if not company_name and not address and not website and not logo_path:
        return  # nothing to show

    # Spacer before footer
    doc.add_paragraph()

    # Thin decorative line
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run("—" * 40)
    _set_run_font(run_line, style["body_font_name"], Pt(8))
    p_line.paragraph_format.space_after = Pt(6)

    # 3-column table
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True

    # Remove all borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn("w:val"), "nil")
        tblBorders.append(border_el)
    tblPr.append(tblBorders)

    # ── Left cell: logo + company name ──
    left_cell = table.cell(0, 0)
    # Remove cell borders
    _remove_cell_borders(left_cell)

    if logo_path:
        logo_full = _resolve_relative_path(logo_path)
        if logo_full and logo_full.exists():
            try:
                pil_logo = PILImage.open(str(logo_full))
                if pil_logo.mode not in ("RGB", "L"):
                    pil_logo = pil_logo.convert("RGB")
                # Scale logo to ~0.8 cm height
                logo_h_px = int(0.8 / Cm(1) * pil_logo.height) if pil_logo.height > 0 else pil_logo.height
                ratio = 0.8 / (pil_logo.height / 37.795) if pil_logo.height > 0 else 1
                new_w = int(pil_logo.width * ratio) if ratio < 1 else pil_logo.width
                if ratio < 1:
                    pil_logo = pil_logo.resize((new_w, int(pil_logo.height * ratio)), PILImage.LANCZOS)
                buf = std_io.BytesIO()
                pil_logo.save(buf, format="PNG")
                buf.seek(0)
                logo_p = left_cell.paragraphs[0]
                logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                logo_run = logo_p.add_run()
                logo_run.add_picture(buf, width=Cm(1.2))
                # Company name next to logo
                name_run = logo_p.add_run(f"  {company_name}")
                _set_run_font(name_run, style["body_font_name"], Pt(10))
            except Exception:
                _set_company_info_text(left_cell.paragraphs[0], company_name, style)
        else:
            _set_company_info_text(left_cell.paragraphs[0], company_name, style)
    else:
        _set_company_info_text(left_cell.paragraphs[0], company_name, style)

    # ── Center cell: address ──
    center_cell = table.cell(0, 1)
    _remove_cell_borders(center_cell)
    cp = center_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if address:
        addr_run = cp.add_run(f"地址：{address}")
        _set_run_font(addr_run, style["body_font_name"], Pt(9))

    # ── Right cell: website ──
    right_cell = table.cell(0, 2)
    _remove_cell_borders(right_cell)
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if website:
        web_run = rp.add_run(website)
        _set_run_font(web_run, style["body_font_name"], Pt(9))
        web_run.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)  # link blue


def _set_company_info_text(paragraph, text, style):
    """Helper: set centred company name text in a paragraph."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    _set_run_font(run, style["body_font_name"], Pt(10))


def _remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn("w:val"), "nil")
        tcBorders.append(border_el)
    tcPr.append(tcBorders)


def _add_attachments_section(doc, attachments, style):
    """Add a "资质证书附件" section with all available images.

    Parameters
    ----------
    attachments : list[dict]
        Each dict: {"path": str, "label": str, "type": str (optional)}
        path — absolute or relative filesystem path to the image
        label — description for the document (e.g. "营业执照")
        type — "id_card" for ID cards (no rotation, placed side-by-side in pairs)
    """
    if not attachments:
        return

    # ── Section heading ──
    h = doc.add_heading("附件：证照及相关文件", level=1)
    for run in h.runs:
        _set_run_font(
            run,
            style['heading1_font_name'],
            style['heading1_font_size'],
            bold=True,
        )

    i = 0
    while i < len(attachments):
        att = attachments[i]
        path = att.get('path', '')
        label = att.get('label', '附件')
        att_type = att.get('type', '')

        if not path:
            i += 1
            continue

        # ── ID card pair: place two side-by-side ──
        if att_type == 'id_card' and i + 1 < len(attachments) and attachments[i + 1].get('type') == 'id_card':
            next_att = attachments[i + 1]
            next_path = next_att.get('path', '')
            next_label = next_att.get('label', '附件')

            # Half-width for each image to fit two side by side
            half_width = Cm(6.8)

            # Build in-memory images for both ID cards
            img_bufs = []
            img_labels = []
            for img_path, img_label in [(path, label), (next_path, next_label)]:
                fp = _resolve_relative_path(img_path)
                if fp is None:
                    img_bufs.append(None)
                    img_labels.append(img_label)
                    continue
                try:
                    pil_img = PILImage.open(str(fp))
                except Exception:
                    img_bufs.append(None)
                    img_labels.append(img_label)
                    continue
                if pil_img.mode not in ('RGB', 'L'):
                    pil_img = pil_img.convert('RGB')
                # No rotation for ID cards
                max_px = int(half_width / Cm(1) * 37.795)
                if pil_img.width > max_px:
                    ratio = max_px / pil_img.width
                    pil_img = pil_img.resize((max_px, int(pil_img.height * ratio)), PILImage.LANCZOS)
                buf = std_io.BytesIO()
                pil_img.save(buf, format='PNG')
                buf.seek(0)
                img_bufs.append(buf)
                img_labels.append(img_label)

            # Create a 1-row 2-column borderless table
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True

            for col_idx, (buf, lbl) in enumerate(zip(img_bufs, img_labels)):
                cell = table.cell(0, col_idx)
                # Remove cell borders
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ('top', 'left', 'bottom', 'right'):
                    border_el = OxmlElement(f'w:{border_name}')
                    border_el.set(qn('w:val'), 'nil')
                    tcBorders.append(border_el)
                tcPr.append(tcBorders)

                if buf is None:
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(f'[图片缺失: {lbl}]')
                    _set_run_font(run, style['body_font_name'], Pt(10))
                else:
                    # Get image width from buffer
                    buf.seek(0)
                    check = PILImage.open(buf)
                    width_in = check.width / 96.0
                    buf.seek(0)

                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(buf, width=Inches(min(width_in, 2.7)))

                    # Label below image in cell
                    p2 = cell.add_paragraph(lbl)
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r2 = p2.add_run(lbl)
                    _set_run_font(r2, style['body_font_name'], Pt(10))

            # Remove table borders
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border_el = OxmlElement(f'w:{border_name}')
                border_el.set(qn('w:val'), 'nil')
                tblBorders.append(border_el)
            tblPr.append(tblBorders)

            i += 2
            continue

        # ── Regular (non-ID-card) attachment ──
        _insert_image(doc, path, label, style, no_rotate=(att_type == 'id_card'))
        i += 1

    # spacer after all attachments
    doc.add_paragraph()


# ── Public API ────────────────────────────────────────────────────────

def render_bid_to_docx(chapters, project_name, style_config=None, chapter_images=None, company_name=""):
    """Render bid chapters into a formatted ``.docx`` file.

    Parameters
    ----------
    chapters : list[dict]
        Each dict must have ``"title"`` (str) and ``"content"`` (str).
    project_name : str
        Used on the cover page and in the output filename.
    style_config : dict | None
        Optional overrides merged on top of ``DEFAULT_STYLE``.
    chapter_images : list[list[dict]] | None
        Optional per-chapter images to insert inline. Same length as chapters.
        Each inner list contains ``{"path": str, "label": str}`` dicts.
    company_name : str
        Company name for cover page and header. Pulled from CompanyProfile.

    Returns
    -------
    str
        Absolute path to the generated ``.docx`` file.
    """
    style = dict(DEFAULT_STYLE)
    if style_config:
        style.update(style_config)
    # Use company name as header_text default if not explicitly configured
    # Also strip any legacy hardcoded default that may still exist in old templates
    if not style.get("header_text") or style.get("header_text") == "云南宏曦科技有限公司":
        style["header_text"] = company_name

    doc = Document()

    # ── Page setup (default section for the whole document) ──
    section = doc.sections[0]
    section.top_margin = style["margin_top"]
    section.bottom_margin = style["margin_bottom"]
    section.left_margin = style["margin_left"]
    section.right_margin = style["margin_right"]

    # ── Header ──
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_run = hp.add_run(style["header_text"])
    _set_run_font(h_run, style["heading1_font_name"], Pt(9))

    # ── Footer: company info line + page number ──
    footer = section.footer
    footer.is_linked_to_previous = False

    # Load company profile for footer
    footer_company = ""
    footer_address = ""
    try:
        from app.models.company_profile import CompanyProfile
        from sqlalchemy import create_engine, select
        from sqlalchemy.orm import Session
        sync_url = settings.DATABASE_URL_SYNC
        engine = create_engine(sync_url)
        with Session(engine) as session:
            result = session.execute(select(CompanyProfile).limit(1))
            profile = result.scalar_one_or_none()
            if profile:
                footer_company = profile.company_name or ""
                footer_address = profile.address or ""
    except Exception:
        pass

    # Company info line (name left, address right) — borderless 2-col table
    if footer_company or footer_address:
        ft = footer.add_table(rows=1, cols=2, width=Cm(16))
        ft.autofit = True
        # Remove table borders
        tbl = ft._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblBorders = OxmlElement("w:tblBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border_el = OxmlElement(f"w:{border_name}")
            border_el.set(qn("w:val"), "nil")
            tblBorders.append(border_el)
        tblPr.append(tblBorders)
        # Remove cell borders
        for cell in (ft.cell(0, 0), ft.cell(0, 1)):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for bn in ("top", "left", "bottom", "right"):
                be = OxmlElement(f"w:{bn}")
                be.set(qn("w:val"), "nil")
                tcBorders.append(be)
            tcPr.append(tcBorders)
        # Left: company name
        lp = ft.cell(0, 0).paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = lp.add_run(footer_company)
        _set_run_font(lr, style["body_font_name"], Pt(9))
        # Right: address
        rp = ft.cell(0, 1).paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if footer_address:
            rr = rp.add_run(f"地址：{footer_address}")
            _set_run_font(rr, style["body_font_name"], Pt(9))

    # Page number below company info
    fp = footer.add_paragraph()
    _add_page_number(fp)

    # ── Cover page ──
    _add_cover_page(doc, project_name, style, company_name)

    # ── Table of Contents ──
    _add_toc_page(doc, chapters, style)
    _insert_page_break(doc)  # chapters start after TOC

    # ── Body: each chapter ──
    for i, chapter in enumerate(chapters):
        # Insert page break before every chapter except the first
        # (first chapter already follows the TOC page break).
        # Using _insert_page_break() ensures each major section starts
        # on a fresh page and never bleeds from the previous section.
        if i > 0:
            _insert_page_break(doc)
        # Chapter heading – 黑体 16pt (三号) bold
        h = doc.add_heading(chapter["title"], level=1)
        for run in h.runs:
            _set_run_font(
                run,
                style["heading1_font_name"],
                style["heading1_font_size"],
                bold=True,
            )

        # Preprocess content to remove obvious markdown artifacts
        content = chapter.get("content", "")
        content = _preprocess_content(content)

        # Parse content line-by-line, accumulating table rows
        lines = content.split('\n')
        table_rows = []       # accumulates rows for pipe tables
        table_header = True   # first row of a pipe table block is header

        # Semicolon key:value table state
        sc_table_rows = []    # accumulates rows for semicolon tables
        sc_table_keys = []    # keys extracted from the first row
        in_sc_table = False   # currently inside a semicolon table block

        idx = 0
        while idx < len(lines):
            line = lines[idx]
            stripped = line.strip()

            # --- Empty line ---
            if not stripped:
                # Flush any pending pipe table
                if table_rows:
                    _render_table(doc, table_rows, style)
                    table_rows = []
                    table_header = True
                # Flush any pending semicolon table
                if sc_table_rows:
                    _render_table(doc, sc_table_rows, style)
                    sc_table_rows = []
                    sc_table_keys = []
                    in_sc_table = False
                doc.add_paragraph()
                idx += 1
                continue

            # --- Semicolon table title (render as centred bold body text) ---
            if _is_semicolon_table_title(stripped):
                # Flush any pending tables first
                if table_rows:
                    _render_table(doc, table_rows, style)
                    table_rows = []
                    table_header = True
                if sc_table_rows:
                    _render_table(doc, sc_table_rows, style)
                    sc_table_rows = []
                    sc_table_keys = []
                    in_sc_table = False
                # Render table title as body paragraph (centred, bold)
                p_title = doc.add_paragraph()
                _set_paragraph_spacing(p_title, style["body_line_spacing"])
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_title = p_title.add_run(stripped)
                _set_run_font(run_title, style["body_font_name"], style["body_font_size"], bold=True)
                in_sc_table = True
                sc_table_rows = []
                sc_table_keys = []
                idx += 1
                continue

            # --- Semicolon table row (when inside a semicolon table block) ---
            if in_sc_table and _is_semicolon_table_row(stripped):
                vals = _parse_semicolon_table_row(stripped)
                if not sc_table_keys:
                    # First row: extract keys as header
                    sc_table_keys = _extract_semicolon_keys(stripped)
                    sc_table_rows.append(sc_table_keys)
                sc_table_rows.append(vals)
                idx += 1
                continue

            # --- Not in semicolon table anymore if we reach here ---
            if in_sc_table and not _is_semicolon_table_row(stripped):
                # End of semicolon table block
                if sc_table_rows:
                    _render_table(doc, sc_table_rows, style)
                sc_table_rows = []
                sc_table_keys = []
                in_sc_table = False

            # --- Table row (accumulate, don't render yet) ---
            if _is_table_row(stripped):
                table_rows.append(_parse_table_cells(stripped))
                idx += 1
                continue

            # --- Table separator line (skip it, mark that we're in a table) ---
            if _is_table_separator(stripped):
                # Flush any previously accumulated rows before separator
                # (shouldn't normally happen — separator comes after header)
                idx += 1
                continue

            # --- Not a table line → flush pending table first ---
            if table_rows:
                _render_table(doc, table_rows, style)
                table_rows = []
                table_header = True
            if sc_table_rows:
                _render_table(doc, sc_table_rows, style)
                sc_table_rows = []
                sc_table_keys = []
                in_sc_table = False

            # --- Inline ID card pair: [IDPAIR:front_path|front_label|back_path|back_label] ---
            if stripped.startswith('[IDPAIR:') and stripped.endswith(']'):
                inner = stripped[8:-1]
                parts = inner.split('|')
                if len(parts) >= 4:
                    _render_id_card_pair(doc, parts[0], parts[1], parts[2], parts[3], style)
                idx += 1
                continue

            # --- Inline image marker: [IMG:path|label] ---
            if stripped.startswith('[IMG:') and stripped.endswith(']'):
                inner = stripped[5:-1]
                parts = inner.split('|', 1)
                img_path = parts[0]
                img_label = parts[1] if len(parts) > 1 else ''
                if img_path:
                    _insert_image(doc, img_path, img_label, style)
                idx += 1
                continue

            # --- Markdown heading (levels 1-3) ---
            if _is_markdown_heading(stripped):
                level, heading_text = _get_heading_level_and_text(stripped)
                heading_text = _clean_lone_symbols(heading_text)
                if heading_text:
                    if level == 1:
                        h_para = doc.add_heading(heading_text, level=1)
                        for run in h_para.runs:
                            _set_run_font(
                                run,
                                style["heading1_font_name"],
                                style["heading1_font_size"],
                                bold=True,
                            )
                    elif level == 2:
                        h_para = doc.add_heading(heading_text, level=2)
                        for run in h_para.runs:
                            _set_run_font(
                                run,
                                style["heading2_font_name"],
                                style["heading2_font_size"],
                                bold=True,
                            )
                    else:
                        h_para = doc.add_heading(heading_text, level=3)
                        for run in h_para.runs:
                            _set_run_font(
                                run,
                                style["heading3_font_name"],
                                style["heading3_font_size"],
                                bold=True,
                            )
                idx += 1
                continue

            # --- Bullet item ---
            if _is_bullet_item(stripped):
                bullet_text = _get_bullet_text(stripped)
                bullet_text = _clean_lone_symbols(bullet_text)
                if bullet_text:
                    p = doc.add_paragraph()
                    _set_paragraph_spacing(p, style["body_line_spacing"])
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    pf = p.paragraph_format
                    pf.left_indent = Cm(1.27)      # ~4 chars
                    pf.first_line_indent = Cm(-0.63)  # hanging indent for bullet
                    # Use • bullet character
                    run = p.add_run(f"● {bullet_text}")
                    _set_run_font(run, style["body_font_name"], style["body_font_size"])
                idx += 1
                continue

            # --- Numbered item ---
            if _is_numbered_item(stripped):
                num_prefix, num_text = _get_numbered_text(stripped)
                num_text = _clean_lone_symbols(num_text)
                if num_text:
                    p = doc.add_paragraph()
                    _set_paragraph_spacing(p, style["body_line_spacing"])
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    pf = p.paragraph_format
                    pf.left_indent = Cm(1.27)
                    pf.first_line_indent = Cm(-0.63)
                    run = p.add_run(f"{num_prefix} {num_text}")
                    _set_run_font(run, style["body_font_name"], style["body_font_size"])
                idx += 1
                continue

            # --- Regular text with possible inline markdown ---
            cleaned = _clean_lone_symbols(stripped)
            if not cleaned:
                idx += 1
                continue

            segments, had_formatting = _clean_markdown_inline(cleaned)
            if had_formatting:
                _add_body_paragraph_with_bold(doc, segments, style)
            else:
                _add_body_paragraph(doc, cleaned, style)
            idx += 1

        # --- End of chapter: flush any remaining tables ---
        if table_rows:
            _render_table(doc, table_rows, style)
        if sc_table_rows:
            _render_table(doc, sc_table_rows, style)
            table_rows = []
            table_header = True

        # Note: page breaks are now inserted BEFORE each new chapter
        # (see _insert_page_break() at the top of the loop), which
        # guarantees each chapter starts on its own page.

        # ── Insert chapter-specific images (qualifications, certificates) ──
        if chapter_images and i < len(chapter_images) and chapter_images[i]:
            # Add a small gap before images
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(6)
            spacer.paragraph_format.space_after = Pt(6)
            for img in chapter_images[i]:
                img_path = img.get("path", "")
                img_label = img.get("label", "")
                if img_path:
                    _insert_image(doc, img_path, img_label, style)

    # Company info is now rendered in the page footer on every page

    # ── Save to OUTPUT_DIR ──
    output_dir = Path(settings.OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)

    uid8 = uuid.uuid4().hex[:8]
    ts = int(time.time())
    file_path = output_dir / f"{_safe_filename(project_name)}_{uid8}_{ts}.docx"

    doc.save(str(file_path))
    return str(file_path.absolute())


def export_to_pdf(docx_path):
    """Convert a ``.docx`` to PDF via LibreOffice headless (best-effort).

    Parameters
    ----------
    docx_path : str
        Path to the ``.docx`` file to convert.

    Returns
    -------
    str | None
        Absolute path to the generated PDF, or ``None`` if conversion failed
        (e.g. LibreOffice is not installed or times out).
    """
    output_dir = str(Path(docx_path).parent)

    try:
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                str(docx_path),
            ],
            check=True,
            timeout=60,
            capture_output=True,
            text=True,
        )
    except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
        return None

    # Locate the generated PDF (name matches the docx stem)
    pdf_path = Path(output_dir) / (Path(docx_path).stem + ".pdf")
    if pdf_path.exists():
        return str(pdf_path.absolute())

    # Fallback: grab any fresh PDF in the output directory
    for f in sorted(Path(output_dir).glob("*.pdf"), key=os.path.getmtime, reverse=True):
        return str(f.absolute())

    return None


def create_default_template():
    """Create a minimal ``default.docx`` template in ``settings.TEMPLATE_DIR``.

    Returns
    -------
    str
        Absolute path to the created template file.
    """
    template_dir = Path(settings.TEMPLATE_DIR)
    template_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Apply standard bid margins
    section = doc.sections[0]
    section.top_margin = DEFAULT_STYLE["margin_top"]
    section.bottom_margin = DEFAULT_STYLE["margin_bottom"]
    section.left_margin = DEFAULT_STYLE["margin_left"]
    section.right_margin = DEFAULT_STYLE["margin_right"]

    # Placeholder heading
    p = doc.add_paragraph("投标文件模板")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    template_path = template_dir / "default.docx"
    doc.save(str(template_path))

    return str(template_path.absolute())


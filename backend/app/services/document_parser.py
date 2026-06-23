import os, re, zipfile
from pathlib import Path


def parse_document(file_path: str) -> str:
    ext = Path(file_path).suffix.lower()
    if ext in (".docx", ".doc", ".wps"):
        return _parse_doc_family(file_path)
    elif ext == ".pdf":
        return _parse_pdf(file_path)
    else:
        raise ValueError(f"Unsupported format: {ext}")


def _parse_doc_family(file_path: str) -> str:
    # 1. python-docx (handles .docx and ZIP-based .doc/.wps)
    try:
        return _parse_docx(file_path)
    except Exception:
        pass

    # 2. antiword (handles legacy binary .doc)
    try:
        import subprocess
        r = subprocess.run(["antiword", file_path], capture_output=True, text=True, timeout=30)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout.strip()
    except Exception:
        pass

    # 3. ZIP fallback: extract text from any XML/HTML inside
    if zipfile.is_zipfile(file_path):
        text = _extract_text_from_zip(file_path)
        if text and len(text) > 100:
            return text

    # 4. OLE2 fallback: try to extract raw text strings
    text = _extract_text_from_ole2(file_path)
    if text and len(text) > 100:
        return text

    raise RuntimeError(
        f"Cannot parse: {file_path}. Please convert to .docx format."
    )


def _extract_text_from_zip(file_path: str) -> str:
    """Extract readable text from XML files inside a ZIP (WPS, etc)."""
    parts = []
    try:
        with zipfile.ZipFile(file_path, "r") as z:
            for name in z.namelist():
                if name.endswith((".xml", ".html", ".htm")) and "media" not in name.lower():
                    try:
                        content = z.read(name).decode("utf-8", errors="ignore")
                        # Strip XML/HTML tags, keep text
                        text = re.sub(r"<[^>]+>", " ", content)
                        text = re.sub(r"\s+", " ", text).strip()
                        if len(text) > 50:
                            parts.append(text)
                    except Exception:
                        pass
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_text_from_ole2(file_path: str) -> str:
    """Extract raw text from OLE2 compound document."""
    try:
        import olefile
        ole = olefile.OleFileIO(file_path)
        # Try to read the WordDocument stream
        if ole.exists("WordDocument"):
            data = ole.openstream("WordDocument").read()
            # Extract readable text (UTF-16LE encoded in .doc)
            text = _extract_unicode_strings(data)
            if text:
                return text
        # Try 1Table or 0Table
        for stream_name in ole.listdir():
            if isinstance(stream_name, (list, tuple)):
                name = "/".join(stream_name)
            else:
                name = str(stream_name)
            if "Table" in name:
                try:
                    data = ole.openstream(stream_name).read()
                    text = _extract_unicode_strings(data)
                    if text:
                        parts = [text]
                        break
                except Exception:
                    pass
        ole.close()
        return text if 'text' in dir() and text else ""
    except ImportError:
        pass
    except Exception:
        pass
    return ""


def _extract_unicode_strings(data: bytes) -> str:
    """Extract readable Chinese text from binary data."""
    # Try UTF-16LE decoding
    parts = []
    i = 0
    while i < len(data) - 1:
        # Find sequences of printable Chinese/ASCII chars
        char = data[i:i+2]
        try:
            ch = char.decode("utf-16-le")
            if ch.isprintable() and (ord(ch) > 0x2000 or ch.isascii()):
                parts.append(ch)
                i += 2
                continue
        except Exception:
            pass
        i += 1

    text = "".join(parts)
    # Filter: keep only sequences with enough Chinese characters
    lines = []
    for segment in re.split(r"[\x00-\x1f]{3,}", text):
        segment = segment.strip()
        if len(segment) > 5:
            lines.append(segment)
    return "\n".join(lines)


def _parse_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def _parse_pdf(file_path: str) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
            for table in page.extract_tables():
                if table:
                    for row in table:
                        row_text = " | ".join(str(c or "") for c in row)
                        if row_text.strip():
                            parts.append(row_text)
    return "\n".join(parts)
